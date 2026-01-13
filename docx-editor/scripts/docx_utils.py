#!/usr/bin/env python3
"""
Utility script for reading and editing .docx files.
Uses python-docx library for Word document manipulation.

Usage:
    docx_utils.py read <file>                  - Display document content with paragraph numbers
    docx_utils.py search <file> <text>         - Search for text and show context
    docx_utils.py replace <file> <old> <new> [--output <outfile>]  - Find and replace text
    docx_utils.py replace-para <file> <para_num> <new_text> [--output <outfile>]  - Replace entire paragraph
    docx_utils.py insert-after <file> <para_num> <new_text> [--output <outfile>]  - Insert paragraph after
    docx_utils.py append-to-para <file> <para_num> <text_to_append> [--output <outfile>]  - Append text to paragraph
    docx_utils.py delete-text <file> <text_to_delete> [--output <outfile>]  - Delete specific text
"""

import argparse
import sys
import re
from pathlib import Path

# Add venv to path
venv_path = Path(__file__).parent.parent / "venv" / "lib"
for p in venv_path.glob("python*/site-packages"):
    sys.path.insert(0, str(p))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def normalize_text(text):
    """Normalize special characters to standard equivalents for matching."""
    replacements = {
        '\u2018': "'",   # Left single quote
        '\u2019': "'",   # Right single quote
        '\u201C': '"',   # Left double quote
        '\u201D': '"',   # Right double quote
        '\u2013': '-',   # En dash
        '\u2014': '-',   # Em dash
        '\u2026': '...', # Ellipsis
        '\u00A0': ' ',   # Non-breaking space
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def read_document(filepath):
    """Read and display document content with paragraph numbers."""
    doc = Document(filepath)
    print(f"\n=== Document: {filepath} ===\n")

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            # Show style info if available
            style = para.style.name if para.style else "Normal"
            print(f"[P{i}] ({style})")
            print(f"    {text[:500]}{'...' if len(text) > 500 else ''}")
            print()

    # Also show tables if present
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- Table {t_idx} ---")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip()[:50] for cell in row.cells]
            print(f"  Row {r_idx}: {cells}")

    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    return doc


def search_document(filepath, search_text):
    """Search for text in document and show context."""
    doc = Document(filepath)
    search_lower = search_text.lower()
    search_normalized = normalize_text(search_text.lower())
    found = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text
        text_lower = text.lower()
        text_normalized = normalize_text(text_lower)

        if search_lower in text_lower or search_normalized in text_normalized:
            found.append((i, text))
            print(f"\n[P{i}] Found match:")
            print(f"    {text[:1000]}")

    if not found:
        print(f"No matches found for: {search_text}")
    else:
        print(f"\nTotal matches: {len(found)}")

    return found


def replace_text_in_paragraph(para, old_text, new_text):
    """Replace text in a paragraph while preserving formatting as much as possible."""
    # Check if old_text exists in the paragraph
    full_text = para.text
    normalized_full = normalize_text(full_text)
    normalized_old = normalize_text(old_text)

    if old_text not in full_text and normalized_old not in normalized_full:
        return False

    # For simple replacements, try to do it run by run first
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
        normalized_run = normalize_text(run.text)
        if normalized_old in normalized_run:
            # Replace using normalized matching
            run.text = normalized_run.replace(normalized_old, new_text)
            return True

    # If text spans multiple runs, we need to rebuild
    if old_text in full_text or normalized_old in normalized_full:
        # Get the first run's formatting to preserve
        if para.runs:
            # Store formatting from first run
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.font.bold
            italic = first_run.font.italic

        # Do the replacement
        if old_text in full_text:
            new_full = full_text.replace(old_text, new_text)
        else:
            new_full = normalized_full.replace(normalized_old, new_text)

        # Clear all runs and add new text
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_full
            # Restore formatting
            if font_name:
                para.runs[0].font.name = font_name
            if font_size:
                para.runs[0].font.size = font_size
            if bold is not None:
                para.runs[0].font.bold = bold
            if italic is not None:
                para.runs[0].font.italic = italic
        else:
            para.add_run(new_full)
        return True

    return False


def replace_text(filepath, old_text, new_text, output_path=None):
    """Find and replace text throughout document."""
    doc = Document(filepath)
    count = 0

    for para in doc.paragraphs:
        if replace_text_in_paragraph(para, old_text, new_text):
            count += 1

    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_text_in_paragraph(para, old_text, new_text):
                        count += 1

    if output_path:
        doc.save(output_path)
        print(f"Saved to: {output_path}")
    else:
        doc.save(filepath)
        print(f"Updated: {filepath}")

    print(f"Replacements made: {count}")
    return count


def replace_paragraph(filepath, para_num, new_text, output_path=None):
    """Replace entire paragraph content."""
    doc = Document(filepath)

    if para_num < 0 or para_num >= len(doc.paragraphs):
        print(f"Error: Paragraph {para_num} not found. Document has {len(doc.paragraphs)} paragraphs.")
        return False

    para = doc.paragraphs[para_num]
    old_text = para.text

    # Preserve style
    style = para.style

    # Clear and set new text
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

    para.style = style

    if output_path:
        doc.save(output_path)
        print(f"Saved to: {output_path}")
    else:
        doc.save(filepath)
        print(f"Updated: {filepath}")

    print(f"Replaced paragraph {para_num}")
    print(f"Old: {old_text[:100]}...")
    print(f"New: {new_text[:100]}...")
    return True


def insert_after_paragraph(filepath, para_num, new_text, output_path=None):
    """Insert a new paragraph after the specified paragraph number."""
    doc = Document(filepath)

    if para_num < 0 or para_num >= len(doc.paragraphs):
        print(f"Error: Paragraph {para_num} not found. Document has {len(doc.paragraphs)} paragraphs.")
        return False

    # Get the paragraph to insert after
    para = doc.paragraphs[para_num]

    # Insert new paragraph after this one
    # We need to access the underlying XML
    from docx.oxml.ns import qn

    new_para = doc.add_paragraph(new_text)

    # Move the new paragraph to the right position
    para._element.addnext(new_para._element)

    if output_path:
        doc.save(output_path)
        print(f"Saved to: {output_path}")
    else:
        doc.save(filepath)
        print(f"Updated: {filepath}")

    print(f"Inserted new paragraph after P{para_num}")
    return True


def append_to_paragraph(filepath, para_num, text_to_append, output_path=None):
    """Append text to the end of a paragraph."""
    doc = Document(filepath)

    if para_num < 0 or para_num >= len(doc.paragraphs):
        print(f"Error: Paragraph {para_num} not found. Document has {len(doc.paragraphs)} paragraphs.")
        return False

    para = doc.paragraphs[para_num]
    para.add_run(text_to_append)

    if output_path:
        doc.save(output_path)
        print(f"Saved to: {output_path}")
    else:
        doc.save(filepath)
        print(f"Updated: {filepath}")

    print(f"Appended text to paragraph {para_num}")
    return True


def delete_text(filepath, text_to_delete, output_path=None):
    """Delete specific text from document."""
    return replace_text(filepath, text_to_delete, "", output_path)


def main():
    parser = argparse.ArgumentParser(description="Read and edit .docx files")
    subparsers = parser.add_subparsers(dest="command", help="Commands")

    # Read command
    read_parser = subparsers.add_parser("read", help="Display document content")
    read_parser.add_argument("file", help="Path to .docx file")

    # Search command
    search_parser = subparsers.add_parser("search", help="Search for text")
    search_parser.add_argument("file", help="Path to .docx file")
    search_parser.add_argument("text", help="Text to search for")

    # Replace command
    replace_parser = subparsers.add_parser("replace", help="Find and replace text")
    replace_parser.add_argument("file", help="Path to .docx file")
    replace_parser.add_argument("old", help="Text to find")
    replace_parser.add_argument("new", help="Replacement text")
    replace_parser.add_argument("--output", "-o", help="Output file path")

    # Replace paragraph command
    replace_para_parser = subparsers.add_parser("replace-para", help="Replace entire paragraph")
    replace_para_parser.add_argument("file", help="Path to .docx file")
    replace_para_parser.add_argument("para_num", type=int, help="Paragraph number (from read command)")
    replace_para_parser.add_argument("new_text", help="New paragraph text")
    replace_para_parser.add_argument("--output", "-o", help="Output file path")

    # Insert after command
    insert_parser = subparsers.add_parser("insert-after", help="Insert paragraph after")
    insert_parser.add_argument("file", help="Path to .docx file")
    insert_parser.add_argument("para_num", type=int, help="Paragraph number to insert after")
    insert_parser.add_argument("new_text", help="New paragraph text")
    insert_parser.add_argument("--output", "-o", help="Output file path")

    # Append to paragraph command
    append_parser = subparsers.add_parser("append-to-para", help="Append text to paragraph")
    append_parser.add_argument("file", help="Path to .docx file")
    append_parser.add_argument("para_num", type=int, help="Paragraph number")
    append_parser.add_argument("text", help="Text to append")
    append_parser.add_argument("--output", "-o", help="Output file path")

    # Delete text command
    delete_parser = subparsers.add_parser("delete-text", help="Delete specific text")
    delete_parser.add_argument("file", help="Path to .docx file")
    delete_parser.add_argument("text", help="Text to delete")
    delete_parser.add_argument("--output", "-o", help="Output file path")

    args = parser.parse_args()

    if args.command == "read":
        read_document(args.file)
    elif args.command == "search":
        search_document(args.file, args.text)
    elif args.command == "replace":
        replace_text(args.file, args.old, args.new, args.output)
    elif args.command == "replace-para":
        replace_paragraph(args.file, args.para_num, args.new_text, args.output)
    elif args.command == "insert-after":
        insert_after_paragraph(args.file, args.para_num, args.new_text, args.output)
    elif args.command == "append-to-para":
        append_to_paragraph(args.file, args.para_num, args.text, args.output)
    elif args.command == "delete-text":
        delete_text(args.file, args.text, args.output)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
